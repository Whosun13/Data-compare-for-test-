import streamlit as st
import pandas as pd
from io import BytesIO
from thefuzz import fuzz
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sentence_transformers import SentenceTransformer, util
import numpy as np
import textract  # .doc o‘qish uchun

# ------------------ Modelni keshlash ------------------
@st.cache_resource
def load_model():
    return SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')

model = load_model()

# ------------------ Lug‘atlar ------------------
texts = {
    "uz": {
        "title": "📊 Ma'lumotlarni Taqqoslash Platformasi",
        "upload_db": "1️⃣ Ma'lumotlar bazasini yuklang (.xlsx, .csv, .doc, .docx, .txt)",
        "upload_check": "2️⃣ Tekshiriladigan ma'lumotlarni yuklang yoki kiriting",
        "input_method": "Kiritish usuli",
        "file_upload": "Fayl yuklash",
        "manual_input": "Qo'lda kiritish",
        "load_db": "Bazani yuklash",
        "load_check": "Tekshiriladigan ma'lumotlar",
        "input_area": "Ma'lumotlarni kiriting (vergul yoki yangi qatordan ajratib)",
        "db_loaded": "**Yuklangan ma'lumotlar bazasi:**",
        "input_loaded": "**Tekshiriladigan ma'lumotlar:**",
        "select_column_db": "Bazadagi taqqoslanadigan ustunni tanlang",
        "select_column_input": "Tekshiriladigan fayldagi ustunni tanlang",
        "extra_columns": "Natijada ko'rsatish uchun qo'shimcha ustunlar",
        "similarity_slider": "Shakl o'xshashlik foizini tanlang (%)",
        "semantic_slider": "Ma'no o'xshashlik foizini tanlang (%)",
        "compare_btn": "Taqqoslash",
        "results": "Natijalar",
        "download_csv": "📥 Natijani yuklab olish (.csv)",
        "download_xlsx": "📥 Natijani yuklab olish (.xlsx)",
        "download_docx": "📥 Natijani yuklab olish (.docx)",
        "unsupported_format": "Qo'llab-quvvatlanmaydigan format"
    },
    "ru": {
        "title": "📊 Платформа сравнения данных",
        "upload_db": "1️⃣ Загрузите базу данных (.xlsx, .csv, .doc, .docx, .txt)",
        "upload_check": "2️⃣ Загрузите или введите проверяемые данные",
        "input_method": "Способ ввода",
        "file_upload": "Загрузить файл",
        "manual_input": "Ввести вручную",
        "load_db": "Загрузить базу",
        "load_check": "Проверяемые данные",
        "input_area": "Введите данные (через запятую или новую строку)",
        "db_loaded": "**Загруженная база данных:**",
        "input_loaded": "**Проверяемые данные:**",
        "select_column_db": "Выберите столбец для сравнения в базе",
        "select_column_input": "Выберите столбец во входных данных",
        "extra_columns": "Дополнительные столбцы для отображения в результате",
        "similarity_slider": "Выберите процент сходства по форме (%)",
        "semantic_slider": "Выберите процент сходства по смыслу (%)",
        "compare_btn": "Сравнить",
        "results": "Результаты",
        "download_csv": "📥 Скачать результат (.csv)",
        "download_xlsx": "📥 Скачать результат (.xlsx)",
        "download_docx": "📥 Скачать результат (.docx)",
        "unsupported_format": "Неподдерживаемый формат"
    }
}

# ------------------ Matn normalizatsiyasi ------------------
def normalize_text(s):
    if pd.isna(s):
        return ""
    s = str(s).lower()
    apostrophes = ["’", "‘", "`", "ʻ", "‛", "´", "ˊ", "ʽ", "ʾ", "ʿ"]
    for apos in apostrophes:
        s = s.replace(apos, "'")
    return " ".join(s.split())

# ------------------ DOC/DOCX o‘qish ------------------
def read_doc_or_docx(file):
    if file.name.endswith(".docx"):
        file_bytes = file.read()
        file.seek(0)
        doc = Document(BytesIO(file_bytes))
        if doc.tables:
            tables_data = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    tables_data.append(row_data)
            df = pd.DataFrame(tables_data)
            df.columns = df.iloc[0]
            df = df[1:].reset_index(drop=True)
            return df
        full_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        return pd.DataFrame(full_text, columns=["Data"])
    else:  # .doc fayl
        text = textract.process(file.name).decode("utf-8", errors="ignore")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return pd.DataFrame(lines, columns=["Data"])

# ------------------ Fayl yuklash ------------------
def load_file(file):
    if file.name.endswith(".xlsx"):
        return pd.read_excel(file)
    elif file.name.endswith(".csv"):
        return pd.read_csv(file)
    elif file.name.endswith(".doc") or file.name.endswith(".docx"):
        return read_doc_or_docx(file)
    elif file.name.endswith(".txt"):
        text = file.read().decode("utf-8", errors="ignore")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return pd.DataFrame(lines, columns=["Data"])
    else:
        st.error(current_texts["unsupported_format"])
        return None

# ------------------ Word saqlash ------------------
def df_to_word(df):
    doc = Document()
    doc.add_heading(current_texts["results"], level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)
        para = hdr_cells[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        run.font.size = Pt(11)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

# ------------------ Interfeys ------------------
lang = st.selectbox("Til / Язык", options=["O'zbekcha", "Русский"])
current_texts = texts["uz"] if lang == "O'zbekcha" else texts["ru"]

st.title(current_texts["title"])
st.subheader(current_texts["upload_db"])
uploaded_db = st.file_uploader(current_texts["load_db"], type=["xlsx", "csv", "doc", "docx", "txt"])

st.subheader(current_texts["upload_check"])
input_type = st.radio(current_texts["input_method"], [current_texts["file_upload"], current_texts["manual_input"]])

input_data = None
if input_type == current_texts["file_upload"]:
    uploaded_check = st.file_uploader(current_texts["load_check"], type=["xlsx", "csv", "doc", "docx", "txt"])
    if uploaded_check is not None:
        input_data = load_file(uploaded_check)
elif input_type == current_texts["manual_input"]:
    raw_text = st.text_area(current_texts["input_area"])
    if raw_text.strip():
        items = [x.strip() for x in raw_text.replace("\n", ",").split(",") if x.strip()]
        input_data = pd.DataFrame(items, columns=["InputData"])

if uploaded_db is not None:
    df = load_file(uploaded_db)
    if df is not None:
        st.write(current_texts["db_loaded"])
        st.dataframe(df)

        if input_data is not None:
            st.write(current_texts["input_loaded"])
            st.dataframe(input_data)

            column_to_check = st.selectbox(current_texts["select_column_db"], df.columns)
            input_column_to_check = st.selectbox(current_texts["select_column_input"], input_data.columns)
            extra_columns = st.multiselect(current_texts["extra_columns"], [col for col in df.columns if col != column_to_check])

            similarity_threshold = st.slider(current_texts["similarity_slider"], 50, 100, 80, 1)
            semantic_threshold = st.slider(current_texts["semantic_slider"], 50, 100, 80, 1)

if st.button(current_texts["compare_btn"]):
    df["__norm_col__"] = df[column_to_check].apply(normalize_text)
    input_data["__norm_input__"] = input_data[input_column_to_check].apply(normalize_text)

    db_sentences = df[column_to_check].astype(str).unique().tolist()
    input_sentences = input_data[input_column_to_check].astype(str).tolist()

    db_embeddings = model.encode(db_sentences, convert_to_tensor=True)
    input_embeddings = model.encode(input_sentences, convert_to_tensor=True)

    filtered_input = input_data[input_data["__norm_input__"].astype(bool)]

    results = []
    for idx, (original, item) in enumerate(zip(filtered_input[input_column_to_check], filtered_input["__norm_input__"])):
        match_rows = df[df["__norm_col__"] == item]
        exact_match = not match_rows.empty

        similar_items = [val for val in df["__norm_col__"].unique()
                         if fuzz.ratio(item, val) >= similarity_threshold and val != item]

        semantic_similar_items = []
        cos_scores = util.pytorch_cos_sim(input_embeddings[idx], db_embeddings)[0]
        for i, score in enumerate(cos_scores):
            if score >= (semantic_threshold / 100.0) and db_sentences[i] != original:
                semantic_similar_items.append(db_sentences[i])

        extra_data = {}
        for col in extra_columns:
            extra_data[col] = ", ".join(match_rows[col].astype(str).unique()) if exact_match else ""

        results.append({
            "Kiritilgan": original,
            "Mavjud": "Ha" if exact_match else "Yo'q",
            "O'xshashlar": ", ".join(similar_items) if similar_items else "-",
            "Semantik o'xshashlar": ", ".join(set(semantic_similar_items)) if semantic_similar_items else "-",
            **extra_data
        })

    result_df = pd.DataFrame(results)
    st.subheader(current_texts["results"])
    st.dataframe(result_df)

    csv = result_df.to_csv(index=False).encode('utf-8')
    st.download_button(current_texts["download_csv"], csv, "natijalar.csv", "text/csv")

    towrite = BytesIO()
    result_df.to_excel(towrite, index=False, engine='openpyxl')
    towrite.seek(0)
    st.download_button(current_texts["download_xlsx"], towrite, "natijalar.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    word_file = df_to_word(result_df)
    st.download_button(current_texts["download_docx"], word_file, "natijalar.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
